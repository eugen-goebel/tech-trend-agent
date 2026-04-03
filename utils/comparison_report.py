"""
Comparison Report Generator — Side-by-side technology comparison DOCX report.

Takes multiple TechAnalysisResult objects and produces a single report
with comparison tables, strength/limitation matrices, and a combined outlook.
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from agents.analyst import TechAnalysisResult


# ---------------------------------------------------------------------------
# Color palette
# ---------------------------------------------------------------------------
COLOR_DARK = RGBColor(0x1A, 0x37, 0x6C)
COLOR_ACCENT = RGBColor(0x2E, 0x6D, 0xB4)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_LIGHT_BG = "F0F5FF"
COLOR_ALT_BG = "FFFFFF"

# Per-technology accent colors for headers
TECH_COLORS = ["1A6C3C", "1A376C", "6C1A3C", "3C1A6C"]


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_heading(doc: Document, text: str, level: int = 1):
    para = doc.add_heading(text, level=level)
    run = para.runs[0]
    run.font.color.rgb = COLOR_DARK
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = COLOR_ACCENT
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(4)


def _add_hr(doc: Document):
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E6DB4")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _build_cover(doc: Document, technologies: list[str]):
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TECHNOLOGY COMPARISON REPORT")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = COLOR_DARK

    doc.add_paragraph()

    tech_para = doc.add_paragraph()
    tech_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tech_para.add_run(" vs ".join(technologies))
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = COLOR_ACCENT

    doc.add_paragraph()
    doc.add_paragraph()

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    run.font.italic = True

    doc.add_page_break()


def _build_comparison_table(doc: Document, technologies: list[str],
                            analyses: list[TechAnalysisResult], field: str):
    """Build a side-by-side comparison table for a text field."""
    cols = len(technologies)
    table = doc.add_table(rows=2, cols=cols)
    table.style = "Table Grid"
    table.autofit = False

    col_width = Inches(6.0 / cols)
    for col in table.columns:
        for cell in col.cells:
            cell.width = col_width

    for i, tech in enumerate(technologies):
        cell = table.rows[0].cells[i]
        color = TECH_COLORS[i % len(TECH_COLORS)]
        _set_cell_bg(cell, color)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(tech)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_WHITE

    for i, analysis in enumerate(analyses):
        cell = table.rows[1].cells[i]
        _set_cell_bg(cell, COLOR_LIGHT_BG if i % 2 == 0 else COLOR_ALT_BG)
        text = getattr(analysis, field, "")
        para = cell.paragraphs[0]
        run = para.add_run(text)
        run.font.size = Pt(9.5)

    doc.add_paragraph()


def _build_list_comparison(doc: Document, technologies: list[str],
                           analyses: list[TechAnalysisResult], field: str):
    """Build a side-by-side comparison of list fields."""
    cols = len(technologies)
    table = doc.add_table(rows=2, cols=cols)
    table.style = "Table Grid"
    table.autofit = False

    col_width = Inches(6.0 / cols)
    for col in table.columns:
        for cell in col.cells:
            cell.width = col_width

    for i, tech in enumerate(technologies):
        cell = table.rows[0].cells[i]
        color = TECH_COLORS[i % len(TECH_COLORS)]
        _set_cell_bg(cell, color)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(tech)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_WHITE

    for i, analysis in enumerate(analyses):
        cell = table.rows[1].cells[i]
        _set_cell_bg(cell, COLOR_LIGHT_BG if i % 2 == 0 else COLOR_ALT_BG)
        items = getattr(analysis, field, [])
        for item in items:
            para = cell.add_paragraph(style="List Bullet")
            para.paragraph_format.left_indent = Inches(0.15)
            run = para.add_run(item)
            run.font.size = Pt(9)

    doc.add_paragraph()


def _build_key_player_comparison(doc: Document, technologies: list[str],
                                 analyses: list[TechAnalysisResult]):
    """Summary table of top players per technology."""
    headers = ["Technology", "Player", "Focus Area", "Market Position"]
    rows_data = []
    for tech, analysis in zip(technologies, analyses):
        for player in analysis.key_players[:3]:
            rows_data.append((tech, player.name, player.focus_area, player.market_position))

    table = doc.add_table(rows=1 + len(rows_data), cols=4)
    table.style = "Table Grid"
    table.autofit = False
    widths = [Inches(1.3), Inches(1.3), Inches(1.8), Inches(1.8)]
    for i, w in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = w

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_bg(cell, "1A376C")
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_WHITE

    for row_idx, (tech, name, focus, position) in enumerate(rows_data):
        row = table.rows[row_idx + 1]
        bg = COLOR_LIGHT_BG if row_idx % 2 == 0 else COLOR_ALT_BG
        for col_idx, value in enumerate([tech, name, focus, position]):
            cell = row.cells[col_idx]
            _set_cell_bg(cell, bg)
            para = cell.paragraphs[0]
            run = para.add_run(value)
            run.font.size = Pt(9)
            if col_idx <= 1:
                run.font.bold = True

    doc.add_paragraph()


def generate_comparison_report(
    technologies: list[str],
    analyses: list[TechAnalysisResult],
    output_dir: str = "output",
) -> str:
    """
    Generate a side-by-side technology comparison DOCX report.

    Args:
        technologies: List of technology names
        analyses:     Corresponding list of TechAnalysisResult objects
        output_dir:   Output directory

    Returns:
        Absolute path to the generated .docx file
    """
    os.makedirs(output_dir, exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    style = doc.styles["Normal"]
    style.font.size = Pt(10.5)
    style.font.name = "Calibri"

    _build_cover(doc, technologies)

    # Executive Summary comparison
    _add_heading(doc, "Executive Summary Comparison")
    _add_hr(doc)
    _build_comparison_table(doc, technologies, analyses, "executive_summary")

    # Maturity comparison
    _add_heading(doc, "Maturity Assessment")
    _add_hr(doc)
    _build_comparison_table(doc, technologies, analyses, "maturity_assessment")

    # Market Landscape
    _add_heading(doc, "Market Landscape")
    _add_hr(doc)
    _build_comparison_table(doc, technologies, analyses, "market_landscape")

    # Key Players
    _add_heading(doc, "Key Players Comparison")
    _add_hr(doc)
    _build_key_player_comparison(doc, technologies, analyses)

    # Strengths comparison
    _add_heading(doc, "Strengths")
    _add_hr(doc)
    _build_list_comparison(doc, technologies, analyses, "strengths")

    # Limitations comparison
    _add_heading(doc, "Limitations")
    _add_hr(doc)
    _build_list_comparison(doc, technologies, analyses, "limitations")

    # Key Trends
    _add_heading(doc, "Key Trends")
    _add_hr(doc)
    _build_list_comparison(doc, technologies, analyses, "key_trends")

    # Risk Factors
    _add_heading(doc, "Risk Factors")
    _add_hr(doc)
    _build_list_comparison(doc, technologies, analyses, "risk_factors")

    # Future Outlook
    _add_heading(doc, "Future Outlook")
    _add_hr(doc)
    _build_comparison_table(doc, technologies, analyses, "future_outlook")

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        f"Generated by Tech Trend Analysis System  •  {datetime.now().strftime('%Y-%m-%d')}"
    )
    run.font.size = Pt(8.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    safe = "_vs_".join(t.lower().replace(" ", "_")[:20] for t in technologies)
    filename = f"tech_comparison_{safe}_{datetime.now().strftime('%Y%m%d')}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)

    return os.path.abspath(filepath)
