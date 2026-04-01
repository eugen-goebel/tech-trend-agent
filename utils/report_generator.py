"""
Report Generator — Produces a professional DOCX technology trend report.

Transforms the structured TechAnalysisResult from the AnalysisAgent into a
formatted Word document with cover page, key player table, use case matrix, etc.
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from agents.analyst import TechAnalysisResult
from utils.chart_generator import create_key_players_chart, create_use_case_impact_chart


# ---------------------------------------------------------------------------
# Color palette (tech green theme)
# ---------------------------------------------------------------------------
COLOR_DARK_GREEN   = RGBColor(0x1A, 0x6C, 0x3C)   # #1A6C3C — headers
COLOR_MED_GREEN    = RGBColor(0x2E, 0xB4, 0x6D)   # #2EB46D — accents
COLOR_LIGHT_BG     = RGBColor(0xF0, 0xFF, 0xF5)   # #F0FFF5 — table backgrounds
COLOR_STRENGTH_GRN = RGBColor(0x1D, 0x7A, 0x4A)   # strengths
COLOR_LIMIT_RED    = RGBColor(0x9B, 0x1C, 0x1C)   # limitations
COLOR_DRIVER_TEAL  = RGBColor(0x0D, 0x66, 0x6B)   # adoption drivers
COLOR_BARRIER_AMB  = RGBColor(0x92, 0x60, 0x0D)   # adoption barriers
COLOR_WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_LIGHT_GRAY   = RGBColor(0xF7, 0xF8, 0xFA)
COLOR_IMPACT_HIGH  = "C62828"
COLOR_IMPACT_MED   = "E65100"
COLOR_IMPACT_LOW   = "2E7D32"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _set_cell_background(cell, hex_color: str):
    """Set a table cell's background color via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_heading(doc: Document, text: str, level: int = 1):
    """Add a styled section heading."""
    para = doc.add_heading(text, level=level)
    run = para.runs[0]
    run.font.color.rgb = COLOR_DARK_GREEN
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = COLOR_MED_GREEN
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(4)
    return para


def _add_bullet_list(doc: Document, items: list[str], indent: int = 0):
    """Add a bulleted list of items."""
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.left_indent = Inches(0.25 * (indent + 1))
        run = para.add_run(item)
        run.font.size = Pt(10.5)


def _add_horizontal_rule(doc: Document):
    """Add a thin horizontal divider."""
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2EB46D")
    pBdr.append(bottom)
    pPr.append(pBdr)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(6)


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------

def _build_cover_page(doc: Document, technology: str):
    """Create a styled cover page."""
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TECHNOLOGY TREND REPORT")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = COLOR_DARK_GREEN

    doc.add_paragraph()

    tech_para = doc.add_paragraph()
    tech_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = tech_para.add_run(technology)
    run2.font.size = Pt(22)
    run2.font.bold = True
    run2.font.color.rgb = COLOR_MED_GREEN

    doc.add_paragraph()
    doc.add_paragraph()

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    run3.font.size = Pt(11)
    run3.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    run3.font.italic = True

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = sub_para.add_run("AI-Powered Technology Intelligence  •  Multi-Agent Analysis System")
    run4.font.size = Pt(10)
    run4.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    run4.font.italic = True

    doc.add_page_break()


def _build_two_column_table(doc: Document, left_label, left_items, left_color,
                            right_label, right_items, right_color,
                            left_bg, right_bg):
    """Build a two-column comparison table (strengths/limitations, drivers/barriers)."""
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.autofit = False

    col_width = Inches(3.1)
    for col in table.columns:
        for cell in col.cells:
            cell.width = col_width

    # Header row
    for i, (label, color) in enumerate([(left_label, left_color), (right_label, right_color)]):
        cell = table.rows[0].cells[i]
        _set_cell_background(cell, color)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(label)
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_WHITE

    # Content row
    for i, (items, bg) in enumerate([(left_items, left_bg), (right_items, right_bg)]):
        cell = table.rows[1].cells[i]
        _set_cell_background(cell, bg)
        for item in items:
            para = cell.add_paragraph(style="List Bullet")
            para.paragraph_format.left_indent = Inches(0.15)
            run = para.add_run(item)
            run.font.size = Pt(10)


def _build_key_player_table(doc: Document, players):
    """Build a key player comparison table."""
    headers = ["Company", "Description", "Focus Area", "Market Position"]
    table = doc.add_table(rows=1 + len(players), cols=4)
    table.style = "Table Grid"
    table.autofit = False

    widths = [Inches(1.2), Inches(2.3), Inches(1.5), Inches(1.5)]
    for i, width in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = width

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_background(cell, "1A6C3C")
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_WHITE

    # Data rows
    for row_idx, player in enumerate(players):
        row = table.rows[row_idx + 1]
        bg = "F0FFF5" if row_idx % 2 == 0 else "FFFFFF"

        data = [player.name, player.description, player.focus_area, player.market_position]
        for col_idx, value in enumerate(data):
            cell = row.cells[col_idx]
            _set_cell_background(cell, bg)
            para = cell.paragraphs[0]
            run = para.add_run(value)
            run.font.size = Pt(9.5)
            if col_idx == 0:
                run.font.bold = True


def _build_use_case_table(doc: Document, use_cases):
    """Build a use case table with color-coded impact levels."""
    headers = ["Use Case", "Description", "Industry", "Impact"]
    table = doc.add_table(rows=1 + len(use_cases), cols=4)
    table.style = "Table Grid"
    table.autofit = False

    widths = [Inches(1.3), Inches(2.7), Inches(1.2), Inches(0.8)]
    for i, width in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = width

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_background(cell, "1A6C3C")
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_WHITE

    # Data rows
    impact_colors = {"High": COLOR_IMPACT_HIGH, "Medium": COLOR_IMPACT_MED, "Low": COLOR_IMPACT_LOW}

    for row_idx, uc in enumerate(use_cases):
        row = table.rows[row_idx + 1]
        bg = "F0FFF5" if row_idx % 2 == 0 else "FFFFFF"

        data = [uc.title, uc.description, uc.industry, uc.impact_level]
        for col_idx, value in enumerate(data):
            cell = row.cells[col_idx]
            _set_cell_background(cell, bg)
            para = cell.paragraphs[0]
            run = para.add_run(value)
            run.font.size = Pt(9.5)
            if col_idx == 0:
                run.font.bold = True
            if col_idx == 3:
                # Color-code impact level
                impact_bg = impact_colors.get(value, "FFFFFF")
                _set_cell_background(cell, impact_bg)
                run.font.color.rgb = COLOR_WHITE
                run.font.bold = True
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def generate_docx_report(
    technology: str,
    analysis: TechAnalysisResult,
    output_dir: str = "output",
) -> str:
    """
    Generate a professional DOCX technology trend report.

    Args:
        technology:  Technology name used as title
        analysis:    Structured TechAnalysisResult from the AnalysisAgent
        output_dir:  Directory where the file will be saved

    Returns:
        Absolute path to the generated .docx file
    """
    os.makedirs(output_dir, exist_ok=True)

    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

    # Default paragraph style
    style = doc.styles["Normal"]
    style.font.size = Pt(10.5)
    style.font.name = "Calibri"

    # 1. Cover Page
    _build_cover_page(doc, technology)

    # 2. Executive Summary
    _add_heading(doc, "Executive Summary")
    _add_horizontal_rule(doc)
    para = doc.add_paragraph(analysis.executive_summary)
    para.paragraph_format.space_after = Pt(12)

    # 3. Technology Overview
    _add_heading(doc, "Technology Overview")
    _add_horizontal_rule(doc)
    doc.add_paragraph(analysis.technology_overview)

    # 4. Maturity Assessment
    _add_heading(doc, "Maturity Assessment")
    _add_horizontal_rule(doc)
    doc.add_paragraph(analysis.maturity_assessment)

    # 5. Market Landscape
    _add_heading(doc, "Market Landscape")
    _add_horizontal_rule(doc)
    doc.add_paragraph(analysis.market_landscape)

    # 6. Key Players
    _add_heading(doc, "Key Players")
    _add_horizontal_rule(doc)
    _build_key_player_table(doc, analysis.key_players)
    if analysis.key_players:
        doc.add_paragraph()
        chart_buf = create_key_players_chart(analysis)
        doc.add_picture(chart_buf, width=Inches(5.2))
    doc.add_paragraph()

    # 7. Use Cases
    _add_heading(doc, "Use Cases")
    _add_horizontal_rule(doc)
    _build_use_case_table(doc, analysis.use_cases)
    if analysis.use_cases:
        doc.add_paragraph()
        chart_buf = create_use_case_impact_chart(analysis)
        doc.add_picture(chart_buf, width=Inches(3.8))
    doc.add_paragraph()

    # 8. Strengths & Limitations
    _add_heading(doc, "Strengths & Limitations")
    _add_horizontal_rule(doc)
    _build_two_column_table(
        doc,
        "STRENGTHS", analysis.strengths, "1D7A4A",
        "LIMITATIONS", analysis.limitations, "9B1C1C",
        "F0FFF4", "FFF5F5",
    )
    doc.add_paragraph()

    # 9. Adoption Drivers & Barriers
    _add_heading(doc, "Adoption Drivers & Barriers")
    _add_horizontal_rule(doc)
    _build_two_column_table(
        doc,
        "DRIVERS", analysis.adoption_drivers, "0D666B",
        "BARRIERS", analysis.adoption_barriers, "92600D",
        "F0FFFF", "FFFBF0",
    )
    doc.add_paragraph()

    # 10. Key Trends
    _add_heading(doc, "Key Trends")
    _add_horizontal_rule(doc)
    _add_bullet_list(doc, analysis.key_trends)

    # 11. Future Outlook
    _add_heading(doc, "Future Outlook")
    _add_horizontal_rule(doc)
    doc.add_paragraph(analysis.future_outlook)

    # 12. Risk Factors
    _add_heading(doc, "Risk Factors")
    _add_horizontal_rule(doc)
    _add_bullet_list(doc, analysis.risk_factors)

    # Footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run(
        f"Generated by Tech Trend Analysis System  •  {datetime.now().strftime('%Y-%m-%d')}"
    )
    run.font.size = Pt(8.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    # Save
    safe_name = technology.lower().replace(" ", "_").replace("/", "-")
    filename = f"tech_trend_{safe_name}_{datetime.now().strftime('%Y%m%d')}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)

    return os.path.abspath(filepath)
