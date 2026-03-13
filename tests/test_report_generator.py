"""Tests for DOCX report generation."""

import os
from docx import Document

from agents.mock_data import AI_MOCK
from agents.analyst import TechAnalysisResult, KeyPlayer, UseCase
from utils.report_generator import generate_docx_report


class TestReportGeneration:
    """Test the full DOCX generation pipeline."""

    def test_generates_docx_file(self, tmp_path):
        path = generate_docx_report("Artificial Intelligence", AI_MOCK, output_dir=str(tmp_path))
        assert os.path.exists(path)
        assert path.endswith(".docx")

    def test_filename_contains_technology(self, tmp_path):
        path = generate_docx_report("Artificial Intelligence", AI_MOCK, output_dir=str(tmp_path))
        assert "artificial_intelligence" in os.path.basename(path)

    def test_file_not_empty(self, tmp_path):
        path = generate_docx_report("Blockchain", AI_MOCK, output_dir=str(tmp_path))
        assert os.path.getsize(path) > 1000

    def test_docx_is_valid(self, tmp_path):
        path = generate_docx_report("AI", AI_MOCK, output_dir=str(tmp_path))
        doc = Document(path)
        assert len(doc.paragraphs) > 10

    def test_contains_technology_name(self, tmp_path):
        path = generate_docx_report("Quantum Computing", AI_MOCK, output_dir=str(tmp_path))
        doc = Document(path)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Quantum Computing" in all_text

    def test_contains_all_sections(self, tmp_path):
        path = generate_docx_report("AI", AI_MOCK, output_dir=str(tmp_path))
        doc = Document(path)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        expected_sections = [
            "TECHNOLOGY TREND REPORT",
            "Executive Summary",
            "Technology Overview",
            "Maturity Assessment",
            "Market Landscape",
            "Key Players",
            "Use Cases",
            "Strengths & Limitations",
            "Adoption Drivers & Barriers",
            "Key Trends",
            "Future Outlook",
            "Risk Factors",
        ]
        for section in expected_sections:
            assert section in all_text, f"Missing section: {section}"

    def test_contains_tables(self, tmp_path):
        path = generate_docx_report("AI", AI_MOCK, output_dir=str(tmp_path))
        doc = Document(path)
        # Key players + use cases + strengths/limitations + drivers/barriers = 4 tables
        assert len(doc.tables) >= 4

    def test_key_player_table_rows(self, tmp_path):
        path = generate_docx_report("AI", AI_MOCK, output_dir=str(tmp_path))
        doc = Document(path)
        # First table should be key players
        player_table = doc.tables[0]
        expected_rows = 1 + len(AI_MOCK.key_players)
        assert len(player_table.rows) == expected_rows

    def test_creates_output_directory(self, tmp_path):
        new_dir = str(tmp_path / "sub" / "reports")
        path = generate_docx_report("AI", AI_MOCK, output_dir=new_dir)
        assert os.path.exists(path)

    def test_special_characters_in_name(self, tmp_path):
        path = generate_docx_report("AI/ML & Robotics", AI_MOCK, output_dir=str(tmp_path))
        assert os.path.exists(path)


class TestReportWithMinimalData:
    """Test report generation with edge cases."""

    def _minimal_analysis(self):
        return TechAnalysisResult(
            executive_summary="Summary.",
            technology_overview="Overview.",
            maturity_assessment="Early.",
            market_landscape="Small.",
            key_players=[
                KeyPlayer(name="X", description="D", focus_area="F", market_position="P")
            ],
            use_cases=[
                UseCase(title="T", description="D", industry="I", impact_level="Low")
            ],
            strengths=["S"],
            limitations=["L"],
            adoption_drivers=["D"],
            adoption_barriers=["B"],
            future_outlook="Outlook.",
            key_trends=["Trend"],
            risk_factors=["Risk"],
        )

    def test_minimal_data_generates(self, tmp_path):
        path = generate_docx_report("Test", self._minimal_analysis(), output_dir=str(tmp_path))
        assert os.path.exists(path)

    def test_single_player_single_usecase(self, tmp_path):
        path = generate_docx_report("Test", self._minimal_analysis(), output_dir=str(tmp_path))
        doc = Document(path)
        player_table = doc.tables[0]
        assert len(player_table.rows) == 2  # header + 1
